"""
doc_bridge.py — Universal document parser.
Parses PDF, Excel, CSV, TSV, JSON, DOCX into normalized dicts.
Used by /data-bridge skill and file_auto_parser hook.
"""

import json
from pathlib import Path
from typing import Any

SUPPORTED_EXT = {".pdf", ".xlsx", ".xls", ".csv", ".tsv", ".json", ".docx"}


def parse(path: str) -> dict[str, Any]:
    p = Path(path)
    ext = p.suffix.lower()
    if ext in {".xlsx", ".xls"}:
        return _parse_excel(str(p))
    elif ext == ".csv":
        return _parse_csv(str(p), sep=",")
    elif ext == ".tsv":
        return _parse_csv(str(p), sep="\t")
    elif ext == ".json":
        return _parse_json(str(p))
    elif ext == ".pdf":
        return _parse_pdf(str(p))
    elif ext == ".docx":
        return _parse_docx(str(p))
    return {"_error": f"Unsupported: {ext}", "_path": str(p), "_format": "unknown"}


def _parse_excel(path: str) -> dict:
    try:
        import pandas as pd

        sheets = pd.read_excel(path, sheet_name=None)
        result: dict[str, Any] = {
            "_format": "excel",
            "_path": path,
            "_sheets": list(sheets.keys()),
            "data": {},
        }
        for name, df in sheets.items():
            df.columns = [str(c) for c in df.columns]
            result["data"][name] = {
                "columns": list(df.columns),
                "rows": len(df),
                "records": df.where(df.notna(), None).to_dict(orient="records"),
            }
        return result
    except ImportError:
        return {
            "_error": "pandas/openpyxl not installed: pip install pandas openpyxl",
            "_format": "excel",
            "_path": path,
        }


def _parse_csv(path: str, sep: str = ",") -> dict:
    try:
        import pandas as pd

        df = pd.read_csv(path, sep=sep)
        df.columns = [str(c) for c in df.columns]
        return {
            "_format": "csv" if sep == "," else "tsv",
            "_path": path,
            "columns": list(df.columns),
            "rows": len(df),
            "records": df.where(df.notna(), None).to_dict(orient="records"),
        }
    except ImportError:
        return {
            "_error": "pandas not installed: pip install pandas",
            "_format": "csv",
            "_path": path,
        }


def _parse_json(path: str) -> dict:
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    return {"_format": "json", "_path": path, "data": data}


def _parse_pdf(path: str) -> dict:
    try:
        import pdfplumber

        result: dict[str, Any] = {"_format": "pdf", "_path": path, "tables": [], "text": []}
        with pdfplumber.open(path) as pdf:
            result["_pages"] = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                tables = page.extract_tables()
                if tables:
                    for t in tables:
                        result["tables"].append({"page": i + 1, "data": t})
                text = page.extract_text()
                if text and text.strip():
                    result["text"].append({"page": i + 1, "content": text.strip()})
        return result
    except ImportError:
        return {
            "_error": "pdfplumber not installed: pip install pdfplumber",
            "_format": "pdf",
            "_path": path,
        }


def _parse_docx(path: str) -> dict:
    try:
        import docx

        doc = docx.Document(path)
        tables = []
        for table in doc.tables:
            rows = [[cell.text for cell in row.cells] for row in table.rows]
            tables.append(rows)
        return {
            "_format": "docx",
            "_path": path,
            "tables": tables,
            "paragraphs": [p.text for p in doc.paragraphs if p.text.strip()],
        }
    except ImportError:
        return {
            "_error": "python-docx not installed: pip install python-docx",
            "_format": "docx",
            "_path": path,
        }


def diff(a: dict, b: dict, label_a: str = "A", label_b: str = "B") -> dict:
    """Structural diff of two parsed documents."""
    try:
        from deepdiff import DeepDiff

        d = DeepDiff(a, b, ignore_order=True, ignore_numeric_type_changes=True)
        changes = d.get("values_changed", {})
        added = d.get("iterable_item_added", {})
        removed = d.get("iterable_item_removed", {})
        return {
            "summary": {
                f"{label_a}_path": a.get("_path", "?"),
                f"{label_b}_path": b.get("_path", "?"),
                "value_changes": len(changes),
                "added_items": len(added),
                "removed_items": len(removed),
                "total_changes": len(changes) + len(added) + len(removed),
            },
            "details": {
                "values_changed": dict(list(changes.items())[:50]),
                "added": dict(list(added.items())[:20]),
                "removed": dict(list(removed.items())[:20]),
            },
        }
    except ImportError:
        return {"_error": "deepdiff not installed: pip install deepdiff"}


def summarize(parsed: dict) -> str:
    """Single-line human summary of a parsed document."""
    fmt = parsed.get("_format", "unknown")
    name = Path(parsed.get("_path", "?")).name

    if fmt == "excel":
        sheets = parsed.get("_sheets", [])
        total = sum(parsed["data"][s]["rows"] for s in sheets if s in parsed.get("data", {}))
        return f"{name}: Excel {len(sheets)} sheets, {total} total rows"
    elif fmt in ("csv", "tsv"):
        cols = len(parsed.get("columns", []))
        rows = parsed.get("rows", 0)
        return f"{name}: {fmt.upper()} {cols} cols × {rows} rows"
    elif fmt == "json":
        data = parsed.get("data")
        if isinstance(data, list):
            return f"{name}: JSON [{len(data)} items]"
        elif isinstance(data, dict):
            return f"{name}: JSON {{{len(data)} keys}}"
        return f"{name}: JSON"
    elif fmt == "pdf":
        return (
            f"{name}: PDF {parsed.get('_pages', '?')} pages, "
            f"{len(parsed.get('tables', []))} tables, "
            f"{len(parsed.get('text', []))} text blocks"
        )
    elif fmt == "docx":
        return (
            f"{name}: DOCX {len(parsed.get('tables', []))} tables, "
            f"{len(parsed.get('paragraphs', []))} paragraphs"
        )
    elif "_error" in parsed:
        return f"{name}: ERROR — {parsed['_error']}"
    return f"{name}: {fmt}"
